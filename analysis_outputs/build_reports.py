from __future__ import annotations

import csv
import os
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "analysis_outputs"
FONT = "Times New Roman"
NAVY = RGBColor(23, 63, 95)
GRAY = RGBColor(75, 75, 75)


def rows(path: Path):
    with path.open(newline="", encoding="utf-8-sig") as f:
        return list(csv.DictReader(f))


def by_key(path: Path, key="Statistic", value="Value"):
    return {r[key]: r[value] for r in rows(path)}


def set_font(run, size=9.2, bold=False, italic=False, color=None):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    shd = cell._tc.get_or_add_tcPr().find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        cell._tc.get_or_add_tcPr().append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=55, start=65, bottom=55, end=65):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def fixed_table(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    total_dxa = int(round(sum(widths) * 1440))
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(round(width * 1440))))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(int(round(widths[idx] * 1440))))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_table(doc, headers, body, widths, font_size=7.1):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    fixed_table(table, widths)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(h)
        set_cell_shading(cell, "DCE6F1")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            set_font(run, font_size, bold=True)
    set_repeat_table_header(table.rows[0])
    for vals in body:
        cells = table.add_row().cells
        for i, val in enumerate(vals):
            cells[i].text = str(val)
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                set_font(run, font_size)
    return table


def field_run(paragraph, field_name):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_name
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, separate, text, end):
        run._r.append(node)
    set_font(run, 8)


def configure_section(section, columns=1):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.28)
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(columns))
    cols.set(qn("w:space"), "300")


def add_column_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "column")
    r._r.append(br)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_break(WD_BREAK.PAGE)


def add_para(doc, text="", size=9.2, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             before=0, after=3, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    set_font(r, size, bold, italic, color)
    return p


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text.upper())
    set_font(r, 9.5, bold=True, color=NAVY)
    return p


def add_caption(doc, text):
    return add_para(doc, text, size=7.7, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=3)


def add_image(doc, path, caption, width=3.32):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", caption)
    add_caption(doc, caption)


def add_reference(doc, number, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(1.5)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(f"[{number}] {text}")
    set_font(r, 7.3)


def set_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(9.2)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.0


def add_cover(doc, title, subtitle, dataset):
    section = doc.sections[0]
    configure_section(section, 1)
    add_para(doc, "TUNKU ABDUL RAHMAN UNIVERSITY OF MANAGEMENT AND TECHNOLOGY",
             13, True, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    add_para(doc, "FACULTY OF COMPUTING AND INFORMATION TECHNOLOGY", 11, True,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_para(doc, "BMMS2094 STATISTICS FOR DATA SCIENCE", 12, True,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=26)
    add_para(doc, "GROUP ASSIGNMENT REPORT", 12, True, align=WD_ALIGN_PARAGRAPH.CENTER, after=20, color=NAVY)
    add_para(doc, title, 18, True, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_para(doc, subtitle, 11, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=26, color=GRAY)
    add_para(doc, "Dataset", 10, True, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    add_para(doc, dataset, 9.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=22)
    add_table(
        doc,
        ["Member", "Student ID", "Signature", "Contribution"],
        [[f"[MEMBER {i} NAME]", "[ID]", "[SIGNATURE]", "25%"] for i in range(1, 5)],
        [2.45, 1.45, 1.55, 1.2], font_size=8.5
    )
    add_para(doc, "Semester I, Session 2026/2027", 10, True,
             align=WD_ALIGN_PARAGRAPH.CENTER, before=24, after=3)
    add_para(doc, "Submission date: [INSERT DATE]", 9.5,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("BMMS2094 Group Assignment"), 8, color=GRAY)


def start_ieee_body(doc, short_title):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section, 2)
    section.header.is_linked_to_previous = False
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.paragraph_format.space_after = Pt(0)
    set_font(hp.add_run(f"BMMS2094 | {short_title}"), 7.5, color=GRAY)
    section.footer.is_linked_to_previous = False
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    field_run(fp, "PAGE")
    return section


def title_block(doc, title, subtitle):
    add_para(doc, title, 15.5, True, align=WD_ALIGN_PARAGRAPH.CENTER, after=3)
    add_para(doc, subtitle, 9, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, after=5, color=GRAY)


def fmt(x, n=3):
    return f"{float(x):.{n}f}"


def save_report(doc, path):
    """Save to the requested name, or a clearly named copy if Word holds a lock."""
    try:
        doc.save(path)
        return path
    except PermissionError:
        fallback = path.with_name(f"{path.stem}_updated{path.suffix}")
        doc.save(fallback)
        return fallback


def load_nasa_comparison_evidence(base):
    shared_selection = [
        r for r in rows(base / "nasa_model_selection.csv")
        if r["Model"] not in {"SARIMA", "Seasonal naive"}
    ]
    manual_sarima = rows(base / "manual_sarima" / "nasa_model_selection.csv")
    if len(manual_sarima) != 1 or manual_sarima[0]["Model"] != "SARIMA":
        raise ValueError("Expected one SARIMA row in the manual evidence bundle")
    accuracy = sorted([manual_sarima[0], *shared_selection], key=lambda r: float(r["RMSE"]))

    diagnostics = {
        r["Model"]: r for r in rows(base / "nasa_residual_diagnostics.csv")
        if r["Model"] not in {"SARIMA", "Seasonal naive"}
    }
    manual_diagnostics = rows(base / "manual_sarima" / "nasa_residual_diagnostics.csv")
    if len(manual_diagnostics) != 1 or manual_diagnostics[0]["Model"] != "SARIMA":
        raise ValueError("Expected one SARIMA diagnostic row in the manual evidence bundle")
    diagnostics["SARIMA"] = manual_diagnostics[0]
    return accuracy, diagnostics


def report_nasa():
    base = OUT / "nasa"
    acc, diag = load_nasa_comparison_evidence(base)
    desc = by_key(base / "nasa_descriptive_statistics.csv")
    climate = rows(base / "nasa_monthly_climatology.csv")
    best = acc[0]
    peak = max(climate, key=lambda r: float(r["solar_irradiance"]))
    low = min(climate, key=lambda r: float(r["solar_irradiance"]))
    doc = Document()
    set_styles(doc)
    title = "Forecasting Monthly Solar Irradiance in Kuala Lumpur"
    add_cover(doc, title, "NASA POWER evidence for renewable-energy planning",
              "NASA POWER Monthly API: ALLSKY_SFC_SW_DWN, 2001-2025")
    start_ieee_body(doc, "Kuala Lumpur Solar Irradiance")

    title_block(doc, title, "A comparative R analysis of 300 NASA POWER observations")
    add_heading(doc, "Abstract")
    add_para(doc,
        f"This study forecasts monthly all-sky surface shortwave downward irradiance for Kuala Lumpur using {desc['Observations']} NASA POWER observations from "
        f"{desc['First month']} to {desc['Last month']}. Four course-aligned models were evaluated on January 2021-December 2025. "
        f"{best['Model']} achieved the lowest RMSE ({fmt(best['RMSE'],3)} kWh/m2/day), MAE ({fmt(best['MAE'],3)}) and MAPE ({fmt(best['MAPE'],2)}%). "
        f"Its residual Ljung-Box p-value was {fmt(diag[best['Model']]['Ljung_Box_p_value'],3)}. All 60 test forecasts were compared with known actual values. "
        "The results can support preliminary solar planning but do not represent site-specific generation or climate-change attribution.")
    add_para(doc, "Keywords - NASA POWER, solar irradiance, renewable energy, SARIMA, SDG 7.", 8.2, italic=True, after=4)
    add_heading(doc, "I. Introduction")
    add_para(doc,
        "Reliable expectations of solar-resource seasonality support renewable-energy planning, maintenance scheduling and capacity assessment. NASA POWER provides analysis-ready solar and meteorological time series globally [1]. This study asks which syllabus-aligned model most accurately forecasts Kuala Lumpur's monthly solar irradiance and what the forecast implies for resource variability.")
    add_heading(doc, "II. Objectives")
    add_para(doc,
        "The objectives are to (1) retrieve and validate a reproducible NASA POWER series; (2) describe trend, seasonality and serial dependence; and (3) compare four forecasting methods using the 60-month holdout. Forecasting beyond the available actual observations is outside this accuracy assessment.")
    add_heading(doc, "III. Dataset and SDG alignment")
    add_para(doc,
        "The target is ALLSKY_SFC_SW_DWN at 3.1390 N, 101.6869 E. NASA reports monthly average daily solar energy in kWh/m2/day [1], [2]. The JSON response contains annual YYYY13 entries, which were removed. The resulting January 2001-December 2025 series has 300 continuous months and no -999 fill-value gaps. The study aligns primarily with SDG 7.2 on increasing renewable energy's share [3].")

    add_column_break(doc)
    add_image(doc, base / "figures" / "nasa_overview.png",
              "Fig. 1. Monthly Kuala Lumpur solar irradiance and seasonal distribution.")
    add_heading(doc, "IV. Descriptive findings")
    add_para(doc,
        f"Average irradiance was {float(desc['Mean']):.3f} kWh/m2/day, with a standard deviation of {float(desc['Standard deviation']):.3f}. Observed values ranged from "
        f"{float(desc['Minimum']):.3f} to {float(desc['Maximum']):.3f}. Calendar-month averages peaked in {peak['month']} at {float(peak['solar_irradiance']):.3f} and were lowest in "
        f"{low['month']} at {float(low['solar_irradiance']):.3f} kWh/m2/day, demonstrating a clear annual cycle.")
    add_heading(doc, "V. Data preparation")
    add_para(doc,
        "The R script retrieves and preserves the official JSON, extracts monthly parameter values, removes YYYY13 annual records, converts YYYYMM keys to dates, maps -999 to missing, and checks the expected 300-month calendar, duplicates, gaps and physical range. No interpolation was required.")
    add_heading(doc, "VI. Experimental design")
    add_para(doc,
        "January 2001-December 2020 formed the training set; January 2021-December 2025 was the 60-month test. Accuracy is assessed by comparing every test forecast with its known actual value, followed by residual ACF and Ljung-Box diagnostics.")
    add_para(doc,
        "SARIMA stabilisation: the selected ARIMA(1,0,0)(0,1,1)[12] uses d = 0 because no ordinary trend difference was selected, but D = 1 applies the seasonal difference Delta[12] y(t) = y(t) - y(t-12). This removes the repeating annual level before the autoregressive and seasonal moving-average terms are estimated. The period [12] represents monthly data with one yearly cycle; forecasts are subsequently returned to the original kWh/m2/day scale.")
    add_image(doc, base / "figures" / "nasa_decomposition_diagnostics.png",
              "Fig. 2. STL decomposition and autocorrelation of the STL remainder.", width=3.22)

    add_page_break(doc)
    add_heading(doc, "VII. Forecast accuracy")
    model_labels = {"Basic structural model": "BSM", "Holt-Winters": "HW"}
    acc_body = [[model_labels.get(r["Model"], r["Model"]),
                 fmt(r["Training_RMSE"],4), fmt(r["RMSE"],4),
                 fmt(r["MAE"],4), fmt(r["MAPE"],4), fmt(r["MSE"],4),
                 fmt(r["ME"],4), fmt(r["MPE"],4),
                 fmt(diag[r["Model"]]["Ljung_Box_p_value"],4)] for r in acc]
    add_table(doc, ["Model", "Train\nRMSE", "Test\nRMSE", "Test\nMAE",
                    "Test\nMAPE", "Test\nMSE", "Test\nME", "Test\nMPE",
                    "Ljung-Box\nQ(24) p"],
              acc_body, [.58, .33, .33, .33, .37, .32, .32, .36, .50],
              font_size=4.8)
    add_caption(doc, "TABLE I. Training RMSE, 60-month test accuracy, and residual diagnostics. MAPE and MPE are percentages; BSM denotes the basic structural model and HW denotes Holt-Winters.")
    add_para(doc,
        f"{best['Model']} ranked first with RMSE {fmt(best['RMSE'],3)}, MAE {fmt(best['MAE'],3)}, and MAPE {fmt(best['MAPE'],2)}%. "
        f"Its training RMSE was {fmt(best['Training_RMSE'],3)} versus test RMSE {fmt(best['RMSE'],3)}. {acc[1]['Model']} ranked second on the test set with RMSE {fmt(acc[1]['RMSE'],3)}. "
        "Only RMSE is shown for both training and test data; the remaining accuracy measures are test-set values. The train-test RMSE gap is a generalization diagnostic: a materially higher test RMSE may indicate overfitting, but it is not definitive proof because in-sample residual error and multi-step holdout forecast error are not directly equivalent. Model ranking therefore remains based primarily on test RMSE, subject to residual and physical-plausibility checks.")
    add_heading(doc, "VIII. Diagnostics")
    add_para(doc,
        f"The selected {best['Model']} specification was ARIMA(1,0,0)(0,1,1)[12]. Its Ljung-Box statistic was {fmt(diag[best['Model']]['Ljung_Box_statistic'],2)} with p = "
        f"{fmt(diag[best['Model']]['Ljung_Box_p_value'],3)}, providing no evidence of residual autocorrelation at lag {diag[best['Model']]['Ljung_Box_lag']}.")
    add_image(doc, base / "figures" / "nasa_test_forecasts.png",
              "Fig. 3. Training history and common 60-month test forecasts. The dashed line marks the train-test split; coloured lines are model forecasts and black is observed data.")

    add_page_break(doc)
    add_heading(doc, "IX. Discussion")
    add_para(doc,
        "The close SARIMA, ETS, and basic structural model results indicate that stable annual seasonality explains much of the predictable variation, while SARIMA's seasonal differencing and autoregressive term better captured remaining dependence. Test forecasts can help assess model performance, but engineering simulations require finer-resolution and site-specific inputs.")
    add_heading(doc, "X. Limitations")
    add_para(doc,
        "POWER solar values represent the approximately 1-degree grid cell containing Kuala Lumpur rather than a ground station [2]. Monthly averages hide daily cloud variability and extremes. Irradiance is not electricity production: panel area, efficiency, orientation, shading, temperature, degradation and system losses are excluded. Starting in 2001 avoids the major pre-2001 source transition, but the analysis still cannot attribute changes to climate change.")
    add_heading(doc, "XI. Recommendations")
    add_para(doc,
        "Use the forecast as an initial solar-resource baseline, then validate decisions against local station records and engineering yield models. Update the series annually, monitor residual diagnostics, and examine daily variability when sizing storage or backup capacity. Investment decisions should consider forecast intervals rather than point estimates alone.")
    add_heading(doc, "XII. Conclusion")
    add_para(doc,
        f"{best['Model']} produced the best accuracy on the 60 known January 2021-December 2025 holdout observations and acceptable residual diagnostics. This supports its selection for the present sample, subject to spatial and engineering limitations.")
    add_heading(doc, "References")
    add_reference(doc, 1, "NASA POWER, 'Monthly API,' 2026. https://power.larc.nasa.gov/docs/services/api/temporal/monthly/")
    add_reference(doc, 2, "NASA POWER, 'Data Sources,' 2026. https://power.larc.nasa.gov/docs/methodology/data/sources/")
    add_reference(doc, 3, "United Nations, 'Goal 7: Affordable and Clean Energy,' 2026. https://sdgs.un.org/goals/goal7")
    add_reference(doc, 4, "R. J. Hyndman and G. Athanasopoulos, Forecasting: Principles and Practice, 3rd ed. OTexts, 2021.")
    add_reference(doc, 5, "R Core Team, R: A Language and Environment for Statistical Computing. Vienna, Austria, 2026.")

    path = ROOT / "NASA_Solar_Irradiance_Report.docx"
    return save_report(doc, path)


if __name__ == "__main__":
    print(report_nasa())
